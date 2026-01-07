# ============================================================================
# CV EXPORTER - ENHANCED VERSION with Professional Cover Letter
# ============================================================================

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import os
import tempfile
from typing import List
import re
from datetime import datetime

# Optional: Perfect PDF conversion (pip install docx2pdf)
try:
    from docx2pdf import convert as docx_to_pdf_convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False


class CVExporter:
    """Exports tailored CV: Preserves everything except Skills → replaced with categorized version"""

    @staticmethod
    def extract_existing_skills(doc: Document, skills_start: int, skills_end: int) -> dict:
        """Extract existing skills from the CV (both soft and hard skills)"""
        existing_skills = {
            'soft_skills': [],
            'hard_skills': [],
            'raw_text': []
        }
        
        # Get all text from skills section
        for i in range(skills_start + 1, skills_end):
            para_text = doc.paragraphs[i].text.strip()
            if para_text:
                existing_skills['raw_text'].append(para_text)
        
        # Parse the skills
        for text in existing_skills['raw_text']:
            text_lower = text.lower()
            
            # Check if it's a soft skills line
            if 'soft skill' in text_lower:
                # Extract skills after the label
                skills_part = re.split(r'soft skills?:?', text, flags=re.IGNORECASE)
                if len(skills_part) > 1:
                    skills = [s.strip() for s in re.split(r'[•\-\*,]', skills_part[1]) if s.strip()]
                    existing_skills['soft_skills'].extend(skills)
            
            # Check if it's a hard skills line
            elif 'hard skill' in text_lower:
                skills_part = re.split(r'hard skills?:?', text, flags=re.IGNORECASE)
                if len(skills_part) > 1:
                    skills = [s.strip() for s in re.split(r'[•\-\*,]', skills_part[1]) if s.strip()]
                    existing_skills['hard_skills'].extend(skills)
            
            # If no label, try to extract skills anyway
            else:
                # Extract items separated by • or , or -
                skills = [s.strip() for s in re.split(r'[•\-\*,]', text) if s.strip()]
                # If the line looks like technical skills, add to hard skills
                if any(tech_word in text_lower for tech_word in ['python', 'machine learning', 'docker', 'sql', 'programming']):
                    existing_skills['hard_skills'].extend(skills)
                # Otherwise might be soft skills
                elif any(soft_word in text_lower for soft_word in ['communication', 'teamwork', 'problem', 'creativity', 'leadership']):
                    existing_skills['soft_skills'].extend(skills)
                else:
                    # Default to hard skills if unclear
                    existing_skills['hard_skills'].extend(skills)
        
        return existing_skills

    @staticmethod
    def export_docx(original_file_path: str, tailored_skills: List[str]) -> bytes:
        """Preserve original CV, MERGE new skills with existing skills"""
        if not os.path.exists(original_file_path):
            raise FileNotFoundError(f"File not found: {original_file_path}")

        doc = Document(original_file_path)

        # Find skills section with EXACT matching
        skills_start = -1
        skills_end = len(doc.paragraphs)

        # More precise matching - look for exact section headers
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().upper()
            text_clean = ' '.join(text.split())
            
            if text_clean in ['SKILLS', 'TECHNICAL SKILLS', 'COMPETENCIES', 'EXPERTISE', 
                             'TECHNOLOGIES', 'CORE SKILLS', 'HARD SKILLS', 'PROFESSIONAL SKILLS']:
                skills_start = i
                break

        # If exact match not found, try contains but exclude false positives
        if skills_start == -1:
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip().upper()
                text_clean = ' '.join(text.split())
                
                if 'SKILLS' in text_clean and not any(word in text_clean for word in 
                    ['EXPERIENCE', 'EDUCATION', 'PROFESSIONAL', 'WORK', 'INTERNSHIP', 'PROJECT']):
                    skills_start = i
                    break

        if skills_start != -1:
            # Find end of skills section
            next_section_keywords = ['EXPERIENCE', 'EDUCATION', 'PROJECTS', 'CERTIFICATIONS', 
                                    'LANGUAGES', 'AWARDS', 'PUBLICATIONS', 'WORK EXPERIENCE', 
                                    'PROFESSIONAL EXPERIENCE', 'TRAINING', 'CERTIFICATES']
            
            for i in range(skills_start + 1, len(doc.paragraphs)):
                text = doc.paragraphs[i].text.strip().upper()
                text_clean = ' '.join(text.split())
                
                if any(keyword in text_clean for keyword in next_section_keywords):
                    if len(text_clean) < 50:
                        skills_end = i
                        break

            # Extract existing skills before removing
            existing_skills_data = CVExporter.extract_existing_skills(doc, skills_start, skills_end)

            # Remove old skills content (keep the section header)
            paragraphs_to_remove = []
            for i in range(skills_start + 1, skills_end):
                paragraphs_to_remove.append(i)
            
            for i in reversed(paragraphs_to_remove):
                p = doc.paragraphs[i]._element
                p.getparent().remove(p)

            # Merge: Combine existing hard skills with new tailored skills
            all_hard_skills = existing_skills_data['hard_skills'] + tailored_skills
            
            # Deduplicate (case-insensitive)
            seen = set()
            merged_skills = []
            for skill in all_hard_skills:
                skill_lower = skill.lower().strip()
                if skill_lower not in seen and skill_lower:
                    seen.add(skill_lower)
                    merged_skills.append(skill.strip())

            # Insert position
            insert_position = skills_start + 1
            
            # Add a blank line first
            blank_para = doc.add_paragraph()
            blank_para._element.getparent().remove(blank_para._element)
            doc.paragraphs[skills_start]._element.addnext(blank_para._element)

            # Add soft skills first (if they existed)
            if existing_skills_data['soft_skills']:
                soft_para = doc.add_paragraph()
                soft_run = soft_para.add_run("Soft Skills: ")
                soft_run.bold = True
                soft_run.font.size = Pt(11)
                soft_run.font.color.rgb = RGBColor(44, 62, 80)
                
                soft_skills_text = " • ".join(existing_skills_data['soft_skills'])
                soft_para.add_run(soft_skills_text).font.size = Pt(11)
                
                soft_para._element.getparent().remove(soft_para._element)
                doc.paragraphs[insert_position]._element.addnext(soft_para._element)
                insert_position += 1
                
                # Add blank line
                blank_para = doc.add_paragraph()
                blank_para._element.getparent().remove(blank_para._element)
                doc.paragraphs[insert_position]._element.addnext(blank_para._element)
                insert_position += 1

            # Add hard skills (merged - categorized)
            hard_para = doc.add_paragraph()
            hard_run = hard_para.add_run("Hard Skills: ")
            hard_run.bold = True
            hard_run.font.size = Pt(11)
            hard_run.font.color.rgb = RGBColor(44, 62, 80)
            
            hard_skills_text = " • ".join(merged_skills)
            hard_para.add_run(hard_skills_text).font.size = Pt(11)
            
            hard_para._element.getparent().remove(hard_para._element)
            doc.paragraphs[insert_position]._element.addnext(hard_para._element)
            insert_position += 1
            
            # Add blank line at the end
            blank_para = doc.add_paragraph()
            blank_para._element.getparent().remove(blank_para._element)
            doc.paragraphs[insert_position]._element.addnext(blank_para._element)

        else:
            # Fallback: add at the end
            doc.add_page_break()
            doc.add_heading("TECHNICAL SKILLS", level=1)
            
            p = doc.add_paragraph()
            run = p.add_run("Hard Skills: ")
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(44, 62, 80)
            p.add_run(" • ".join(tailored_skills)).font.size = Pt(11)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def export_pdf_from_docx(docx_bytes: bytes) -> bytes | None:
        """Convert DOCX to PDF with perfect layout preservation"""
        if not DOCX2PDF_AVAILABLE:
            return None

        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
            tmp_docx.write(docx_bytes)
            tmp_docx_path = tmp_docx.name

        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
                tmp_pdf_path = tmp_pdf.name
            docx_to_pdf_convert(tmp_docx_path, tmp_pdf_path)
            with open(tmp_pdf_path, 'rb') as f:
                pdf_bytes = f.read()
            os.unlink(tmp_pdf_path)
        finally:
            os.unlink(tmp_docx_path)

        return pdf_bytes

    @staticmethod
    def export(original_file_path: str, tailored_skills: List[str]) -> dict:
        """Main export: returns DOCX and PDF (if possible)"""
        docx_bytes = CVExporter.export_docx(original_file_path, tailored_skills)
        pdf_bytes = CVExporter.export_pdf_from_docx(docx_bytes)

        return {
            'docx': docx_bytes,
            'pdf': pdf_bytes
        }


# ============================================================================
# ENHANCED MOTIVATION LETTER GENERATOR
# ============================================================================

class MotivationLetterGenerator:
    """
    Generates professional, personalized cover letters following the structure:
    1. Personal header with contact info
    2. Company address
    3. Subject line
    4. Personal introduction with education/background
    5. Why the company resonates
    6. Why you're a strong fit (with concrete examples)
    7. Technical skills and tools
    8. Soft skills and work style
    9. Why this role and why you
    10. Closing with portfolio/GitHub links
    """
    
    @staticmethod
    def generate(
        candidate_name: str,
        candidate_email: str,
        candidate_phone: str = "",
        candidate_address: str = "",
        job_title: str = "",
        company_name: str = "",
        matched_skills: List[str] = None,
        responsibilities: List[str] = None,
        candidate_education: str = "",
        portfolio_url: str = "",
        github_url: str = ""
    ) -> str:
        """
        Generate a professional cover letter with rich detail and personalization
        """
        if matched_skills is None:
            matched_skills = []
        if responsibilities is None:
            responsibilities = []
        
        # Current date
        current_date = datetime.now().strftime("%B %d, %Y")
        
        # Categorize skills
        programming_langs = [s for s in matched_skills if s.lower() in ['python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'r', 'scala', 'go']]
        frameworks = [s for s in matched_skills if s.lower() in ['react', 'angular', 'vue', 'django', 'flask', 'spring boot', 'node.js', 'express']]
        data_tools = [s for s in matched_skills if s.lower() in ['sql', 'postgresql', 'mongodb', 'pandas', 'numpy', 'spark', 'hadoop', 'tableau', 'power bi']]
        cloud_devops = [s for s in matched_skills if s.lower() in ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'jenkins', 'ci/cd']]
        ml_ai = [s for s in matched_skills if s.lower() in ['machine learning', 'deep learning', 'tensorflow', 'pytorch', 'nlp', 'computer vision']]
        
        # Build header
        letter = f"{candidate_name}\n"
        if candidate_address:
            letter += f"{candidate_address}\n"
        letter += f"Email: {candidate_email}\n"
        if candidate_phone:
            letter += f"Mobile: {candidate_phone}\n"
        letter += f"{current_date}\n\n"
        
        # Company address
        letter += f"{company_name}\n"
        letter += "Hiring Team\n"
        if job_title:
            letter += f"{job_title}\n"
        letter += "\n"
        
        # Subject
        subject = f"Application for {job_title}" if job_title else "Application for Position"
        letter += f"Subject: {subject}\n\n"
        
        # Salutation
        letter += f"Dear {company_name} Team,\n\n"
        
        # Opening paragraph - Personal introduction
        letter += f"My name is {candidate_name}"
        if candidate_education:
            letter += f", and {candidate_education}"
        letter += f". I am writing to express my strong interest in the {job_title} position at {company_name}.\n\n"
        
        # Why the company resonates
        letter += f"{company_name}'s mission and commitment to innovation deeply resonate with me. "
        letter += f"I am particularly inspired by your approach to delivering cutting-edge solutions and "
        letter += f"creating meaningful impact in your industry. This philosophy aligns perfectly with my "
        letter += f"passion for leveraging technology to solve real-world problems and drive positive change.\n\n"
        
        # Why you're a strong fit
        letter += "Why I am a strong fit for this role\n\n"
        
        letter += "I bring hands-on experience in creating data-driven solutions, building robust applications, "
        letter += "and delivering projects that combine technical excellence with real business value. "
        
        if ml_ai or data_tools:
            letter += "Throughout my experience, I have developed and deployed machine learning solutions, "
            letter += "designed data pipelines, and transformed complex datasets into actionable insights. "
        
        if frameworks or programming_langs:
            letter += "I have built full-stack applications, integrated APIs, and created user-friendly interfaces "
            letter += "that prioritize both functionality and user experience. "
        
        letter += "This experience has taught me how to bridge the gap between technical complexity and practical, "
        letter += "scalable solutions.\n\n"
        
        # Technical skills section
        if matched_skills:
            letter += "Technical Expertise\n\n"
            
            if programming_langs:
                langs_str = ", ".join(programming_langs[:-1]) + f" and {programming_langs[-1]}" if len(programming_langs) > 1 else programming_langs[0]
                letter += f"I am proficient in {langs_str}, enabling me to tackle diverse technical challenges. "
            
            if frameworks:
                letter += f"I have experience with {', '.join(frameworks[:2])} and other modern frameworks, "
                letter += "allowing me to build scalable and maintainable applications. "
            
            if data_tools:
                letter += f"My expertise in {', '.join(data_tools[:3])} helps me design data pipelines, "
                letter += "perform complex analyses, and create compelling visualizations. "
            
            if cloud_devops:
                letter += f"I work with {', '.join(cloud_devops[:2])} to deploy, monitor, and scale applications "
                letter += "in production environments. "
            
            if ml_ai:
                letter += f"My machine learning experience includes {', '.join(ml_ai[:2])}, "
                letter += "where I've built predictive models and deployed AI-driven features."
            
            letter += "\n\n"
        
        # Soft skills and work style
        letter += "Beyond technical skills, I bring creativity, initiative, and meticulous attention to detail. "
        letter += "I work independently, think in solutions, and value quality in everything I deliver. "
        letter += "My strong communication and organizational abilities enable me to collaborate effectively with "
        letter += "cross-functional teams and translate complex technical concepts into clear, actionable insights.\n\n"
        
        # Responsibilities alignment
        if responsibilities:
            letter += "I am particularly excited about the opportunity to:\n"
            for resp in responsibilities[:3]:
                letter += f"• {resp.capitalize() if resp else 'Contribute to innovative projects'}\n"
            letter += "\n"
        
        # Why this role and why you
        letter += f"Why {company_name} and why me?\n\n"
        letter += f"I am highly motivated to contribute to {company_name}'s mission because it combines my "
        letter += "technical expertise with my genuine passion for innovation and continuous learning. "
        letter += "The opportunity to work on impactful projects, collaborate with talented teams, and grow "
        letter += f"professionally in an environment that values excellence excites me greatly.\n\n"
        
        letter += "I am eager to learn from your experienced team, contribute meaningfully to real projects, "
        letter += "and help drive {company_name}'s success through dedication, technical skill, and innovative thinking.\n\n"
        
        # Closing
        letter += f"I would be honored to join {company_name} as a {job_title} and help make a positive impact. "
        
        if portfolio_url or github_url:
            letter += "You can explore examples of my work and projects "
            if portfolio_url and github_url:
                letter += f"on my portfolio ({portfolio_url}) and GitHub ({github_url})."
            elif portfolio_url:
                letter += f"on my portfolio ({portfolio_url})."
            else:
                letter += f"on my GitHub ({github_url})."
        
        letter += "\n\n"
        letter += "Thank you very much for considering my application. I would be delighted to further discuss "
        letter += "how my skills, creativity, and enthusiasm can support your mission.\n\n"
        
        # Sign-off
        letter += "Warm regards,\n\n"
        letter += f"{candidate_name}\n"
        letter += f"{candidate_email}"
        if candidate_phone:
            letter += f"\n{candidate_phone}"
        
        return letter